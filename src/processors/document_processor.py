"""Multi-format document text extraction."""

from dataclasses import dataclass, field
from pathlib import Path

from src.processors.html_processor import HTMLProcessor


@dataclass
class ProcessedDocument:
    raw_content: str = ""
    clean_text: str = ""
    metadata: dict = field(default_factory=dict)
    attachments: list[str] = field(default_factory=list)


class DocumentProcessor:
    """Extract text from html, pdf, docx, and plain text files."""

    def __init__(self):
        self.html_processor = HTMLProcessor()

    def process_path(self, path: str | Path) -> ProcessedDocument:
        file_path = Path(path)
        suffix = file_path.suffix.lower()
        if suffix in [".html", ".htm"]:
            raw = file_path.read_text(encoding="utf-8")
            return ProcessedDocument(raw_content=raw, clean_text=self.html_processor.process(raw), metadata={"format": "html"})
        if suffix == ".txt":
            raw = file_path.read_text(encoding="utf-8")
            return ProcessedDocument(raw_content=raw, clean_text=raw.strip(), metadata={"format": "text"})
        if suffix == ".pdf":
            return self._process_pdf(file_path)
        if suffix == ".docx":
            return self._process_docx(file_path)
        raise ValueError(f"unsupported document format: {suffix}")

    def process_html(self, raw_html: str) -> ProcessedDocument:
        return ProcessedDocument(raw_content=raw_html, clean_text=self.html_processor.process(raw_html), metadata={"format": "html"})

    def _process_pdf(self, path: Path) -> ProcessedDocument:
        try:
            import pdfplumber
        except ImportError:
            return ProcessedDocument(metadata={"format": "pdf", "error": "pdfplumber not installed"}, attachments=[str(path)])

        parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
        clean_text = "\n".join(parts)
        return ProcessedDocument(clean_text=clean_text, metadata={"format": "pdf", "pages": len(parts)}, attachments=[str(path)])

    def _process_docx(self, path: Path) -> ProcessedDocument:
        try:
            from docx import Document
        except ImportError:
            return ProcessedDocument(metadata={"format": "docx", "error": "python-docx not installed"}, attachments=[str(path)])

        try:
            doc = Document(str(path))
            clean_text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
            return ProcessedDocument(clean_text=clean_text, metadata={"format": "docx"}, attachments=[str(path)])
        except Exception as e:
            return ProcessedDocument(metadata={"format": "docx", "error": str(e)}, attachments=[str(path)])
